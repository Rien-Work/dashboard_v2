from docx import Document

#document maken en titel toevoegen
def document_maken():
    document = Document('test.docx')
    document.save('demo.docx')

def inleiding(type_gebouw, bvo):
    document = Document('demo.docx')
    document.add_heading("Inleiding", 0)
    document.add_paragraph("Het gaat hier om een gebouwtype " + type_gebouw + " met een opervlakte van " + bvo + "m2")
    document.save('demo.docx')

#luchtbehandelingsdeel toevoegen
def algemeen(concept_luchtbehandeling):
    document = Document('demo.docx')
    document.add_heading("Algemeen", 0)
    document.add_paragraph("heeft een luchtbehandeling van " + concept_luchtbehandeling)
    document.save('demo.docx')