import streamlit as st
from docxcompose.composer import Composer
from docx import Document

def document_genereren(title, projectnummer):
    st.balloons()
    st.success('Document gereed, klik op downloaden voor word bestand')
    with open("demo.docx", "rb") as fp:
        btn = st.download_button(
            label="Download bestand",
            data=fp,
            file_name=projectnummer + " " + title + ".docx",
            mime="application/docx"
        )


def tekst_combineren():
    document = Document('demo.docx')
    samen = Document('documenten/punten.docx')
    document.add_paragraph('komt deze test er goed in?')
    document.save('demo.docx')
    master = Document("demo.docx")
    composer = Composer(master)
    doc1 = Document("documenten/punten.docx")
    composer.append(doc1)
    composer.save("demo.docx")


def tekst_combineren_upload():
    document = Document('demo.docx')
    samen = Document('documenten/punten.docx')
    document.add_paragraph('komt deze test er goed in?')
    document.save('demo.docx')
    master = Document("demo.docx")
    composer = Composer(master)
    doc1 = Document("documenten/punten.docx")
    composer.append(doc1)
    composer.save("demo.docx")